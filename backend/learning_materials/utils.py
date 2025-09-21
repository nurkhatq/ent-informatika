# learning_materials/utils.py (запасной вариант с base64)
from docx import Document
import logging
import base64
import uuid

logger = logging.getLogger(__name__)

class DocxToHtmlConverter:
    def __init__(self):
        self.images = []
        self.has_tables = False
        self.has_formulas = False
        self.image_map = {}

    def _style_to_tag(self, style_name: str) -> str:
        if style_name is None:
            return 'p'
            
        if 'heading' in style_name.lower():
            level = ''.join(filter(str.isdigit, style_name)) or '2'
            return f'h{level}'
        return 'p'

    def _process_run(self, run, material_id=None) -> str:
        if not hasattr(run, 'text'):
            return ''
            
        if not run.text.strip() and not hasattr(run, '_element'):
            return ''
            
        if hasattr(run, '_element') and run._element is None:
            return ''

        # Проверяем изображения и сразу встраиваем как base64
        image_html = ''
        if hasattr(run, '_element'):
            try:
                drawings = run._element.findall('.//w:drawing', run._element.nsmap)
                for drawing in drawings:
                    blips = drawing.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                    for blip in blips:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in self.image_map:
                            img_index = self.image_map[embed]
                            if img_index < len(self.images):
                                img_info = self.images[img_index]
                                # Создаем base64 изображение сразу
                                b64_data = base64.b64encode(img_info['image_data']).decode('utf-8')
                                img_src = f"data:image/{img_info['format']};base64,{b64_data}"
                                image_html += f'<img src="{img_src}" alt="Изображение {img_index+1}" class="material-image" style="max-width: 100%; height: auto; margin: 1rem 0;" />'
            except Exception as e:
                logger.error(f"Error  processing drawing: {e}")
        text = run.text
        if not text.strip() and not image_html:
            return ''

        # Экранирование HTML
        text = (text.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
                    .replace('"', '&quot;')
                    .replace("'", '&#39;'))

        # Форматирование
        if hasattr(run, 'bold') and hasattr(run, 'italic'):
            if run.bold and run.italic:
                text = f'<strong><em>{text}</em></strong>'
            elif run.bold:
                text = f'<strong>{text}</strong>'
            elif run.italic:
                text = f'<em>{text}</em>'
            elif hasattr(run, 'underline') and run.underline:
                text = f'<u>{text}</u>'

        return text + image_html

    def _convert_paragraph(self, paragraph, material_id=None) -> str:
        if not hasattr(paragraph, 'text') or not paragraph.text.strip():
            has_images = False
            if hasattr(paragraph, 'runs'):
                for run in paragraph.runs:
                    if hasattr(run, '_element') and run._element is not None:
                        drawings = run._element.findall('.//w:drawing', run._element.nsmap)
                        if drawings:
                            has_images = True
                            break
            
            if not has_images:
                return ''

        if not hasattr(paragraph, 'style') or paragraph.style is None:
            tag = 'p'
        else:
            tag = self._style_to_tag(paragraph.style.name)
            
        if not hasattr(paragraph, 'runs'):
            return f'<{tag}>{paragraph.text}</{tag}>'
            
        content = ''.join(self._process_run(run, material_id) for run in paragraph.runs)
        
        if not content:
            return ''
            
        return f'<{tag}>{content}</{tag}>'

    def _convert_table(self, table, material_id=None) -> str:
        self.has_tables = True
        html = ['<div class="table-wrapper"><table class="border-collapse w-full">']
        
        for i, row in enumerate(table.rows):
            html.append('<tr>')
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'
                content = ''.join(self._convert_paragraph(p, material_id) for p in cell.paragraphs)
                html.append(f'<{tag} class="border p-2">{content}</{tag}>')
            html.append('</tr>')
        
        html.append('</table></div>')
        return ''.join(html)

    def _detect_image_format(self, image_data):
        if image_data.startswith(b'\x89PNG\r\n\x1a\n'):
            return 'png'
        elif image_data.startswith(b'\xff\xd8\xff'):
            return 'jpeg'
        elif image_data.startswith(b'\x47\x49\x46\x38'):
            return 'gif'
        else:
            return 'png'

    def _extract_images(self, document):
        for rel_id, rel in document.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    image_format = self._detect_image_format(image_data)
                    
                    img_index = len(self.images)
                    self.images.append({
                        'image_data': image_data,
                        'format': image_format,
                        'rel_id': rel_id,
                    })
                    
                    self.image_map[rel_id] = img_index
                    
                except Exception as e:
                    logger.error(f"Error extracting image {rel_id}: {str(e)}")

    def convert(self, docx_path, material_id=None):
        try:
            logger.info(f"Converting DOCX: {docx_path}")
            doc = Document(docx_path)
            
            # Извлекаем изображения
            self._extract_images(doc)
            logger.info(f"Extracted {len(self.images)} images")
            
            # Обрабатываем контент
            all_elements = []
            
            try:
                for element in doc.element.body:
                    if element.tag.endswith('p'):
                        for paragraph in doc.paragraphs:
                            if paragraph._element == element:
                                para_html = self._convert_paragraph(paragraph, material_id)
                                if para_html:
                                    all_elements.append(para_html)
                                break
                    elif element.tag.endswith('tbl'):
                        for table in doc.tables:
                            if table._tbl == element:
                                table_html = self._convert_table(table, material_id)
                                all_elements.append(table_html)
                                break
            except Exception as e:
                logger.error(f"Error processing document elements: {e}")
            
            content_html = ''.join(all_elements)
            
            # Для base64 не нужно сохранять отдельные файлы
            saved_images = []
            
            return {
                'html': content_html,
                'images': self.images,
                'saved_images': saved_images,
                'has_tables': self.has_tables,
                'has_formulas': self.has_formulas
            }
            
        except Exception as e:
            logger.exception(f"Error converting document: {e}")
            raise